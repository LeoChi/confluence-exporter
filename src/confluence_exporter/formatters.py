"""Format-specific serializers for exported pages (Strategy pattern).

Each Formatter renders a Confluence page's cleaned HTML to a specific output
format: Markdown, DOCX, PDF, or plain HTML. Formatters are selected by the
:class:`~confluence_exporter.exporter.SpaceExporter` based on the configured
export format.
"""

from __future__ import annotations

import re
from abc import ABC, abstractmethod
from pathlib import Path

from confluence_exporter.logging_utils import get_logger
from confluence_exporter.paths import safe_write_bytes

logger = get_logger()


class Formatter(ABC):
    """Render ``html`` (with page title + breadcrumbs) to a file at ``output_path``."""

    name: str = "abstract"
    extension: str = ""

    @abstractmethod
    def write(
        self,
        *,
        html_body: str,
        output_path: str,
        page_title: str = "",
        breadcrumbs: str = "",
    ) -> None: ...


# ---------------------------------------------------------------------------
# HTML
# ---------------------------------------------------------------------------


class HTMLFormatter(Formatter):
    name = "html"
    extension = "html"

    def write(self, *, html_body, output_path, page_title="", breadcrumbs=""):
        parts = [
            "<!DOCTYPE html>",
            "<html><head><meta charset='utf-8'>",
            f"<title>{page_title}</title>",
            "<style>",
            "body{font-family:Helvetica,Arial,sans-serif;font-size:14px;"
            "max-width:900px;margin:20px auto;padding:0 20px;color:#172B4D;}",
            "h1{font-size:2em;}",
            "table{border-collapse:collapse;margin:10px 0;}",
            "th,td{border:1px solid #DFE1E6;padding:6px 10px;}",
            "th{background:#F4F5F7;}",
            "pre{background:#F4F5F7;padding:10px;overflow-x:auto;}",
            "code{background:#F4F5F7;padding:1px 4px;}",
            ".breadcrumbs{color:#5E6C84;font-size:0.9em;margin-bottom:8px;}",
            "</style></head><body>",
        ]
        if breadcrumbs:
            parts.append(f'<div class="breadcrumbs">{breadcrumbs}</div>')
        if page_title:
            parts.append(f"<h1>{page_title}</h1>")
        parts.append(html_body)
        parts.append("</body></html>")
        safe_write_bytes(output_path, "\n".join(parts).encode("utf-8"))


# ---------------------------------------------------------------------------
# Markdown
# ---------------------------------------------------------------------------


class MarkdownFormatter(Formatter):
    name = "md"
    extension = "md"

    def write(self, *, html_body, output_path, page_title="", breadcrumbs=""):
        from markdownify import markdownify as md

        md_body = md(html_body, heading_style="ATX")
        md_body = re.sub(r"\n{3,}", "\n\n", md_body).strip()

        header = []
        if page_title:
            header.append(f"# {page_title}\n")
        if breadcrumbs:
            header.append(f"> {breadcrumbs}\n")
        out = "\n".join(header) + "\n" + md_body + "\n"
        safe_write_bytes(output_path, out.encode("utf-8"))


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


class DocxFormatter(Formatter):
    name = "docx"
    extension = "docx"

    def write(self, *, html_body, output_path, page_title="", breadcrumbs=""):
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        if page_title:
            heading = doc.add_heading(page_title, level=1)
            for run in heading.runs:
                run.font.size = Pt(18)
        if breadcrumbs:
            p = doc.add_paragraph(breadcrumbs)
            p.runs[0].italic = True

        # Lightweight HTML walk — the full-fidelity converter lives in
        # :mod:`confluence_exporter.converter` and is used by the Convert mode.
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(html_body, "lxml")
        for element in soup.find_all(
            ["h1", "h2", "h3", "h4", "p", "li", "pre", "table"]
        ):
            tag = element.name
            text = element.get_text(strip=True)
            if not text:
                continue
            if tag == "h1":
                doc.add_heading(text, level=2)
            elif tag == "h2":
                doc.add_heading(text, level=3)
            elif tag == "h3":
                doc.add_heading(text, level=4)
            elif tag == "h4":
                doc.add_heading(text, level=5)
            elif tag == "pre":
                p = doc.add_paragraph(text)
                for run in p.runs:
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
            elif tag == "li":
                doc.add_paragraph(text, style="List Bullet")
            elif tag == "table":
                self._add_table(doc, element)
            else:  # <p>
                doc.add_paragraph(text)

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

    @staticmethod
    def _add_table(doc, table_el):
        rows = table_el.find_all("tr")
        if not rows:
            return
        cols = max(len(tr.find_all(["td", "th"])) for tr in rows)
        if cols == 0:
            return
        tbl = doc.add_table(rows=len(rows), cols=cols)
        tbl.style = "Light Grid"
        for i, tr in enumerate(rows):
            cells = tr.find_all(["td", "th"])
            for j in range(cols):
                tbl.rows[i].cells[j].text = (
                    cells[j].get_text(strip=True) if j < len(cells) else ""
                )


# ---------------------------------------------------------------------------
# PDF — uses the configurable engine registry
# ---------------------------------------------------------------------------


class PDFFormatter(Formatter):
    name = "pdf"
    extension = "pdf"

    def __init__(self, engine: str = "auto"):
        self.engine_preference = engine

    def write(self, *, html_body, output_path, page_title="", breadcrumbs=""):
        from confluence_exporter.pdf_engines import render_html_to_pdf

        # Render a tiny wrapper HTML so title + breadcrumbs show up
        doc = []
        doc.append("<!DOCTYPE html><html><head><meta charset='utf-8'>")
        doc.append(f"<title>{page_title}</title>")
        doc.append(
            "<style>"
            "body{font-family:Helvetica,Arial,sans-serif;font-size:10pt;}"
            "h1{font-size:18pt;color:#172B4D;}"
            "table{border-collapse:collapse;margin:8px 0;}"
            "th,td{border:1px solid #DFE1E6;padding:4px 8px;}"
            "th{background:#F4F5F7;}"
            ".breadcrumbs{color:#5E6C84;font-size:9pt;margin-bottom:6pt;}"
            "pre{background:#F4F5F7;padding:8px;font-size:9pt;}"
            "</style></head><body>"
        )
        if breadcrumbs:
            doc.append(f'<div class="breadcrumbs">{breadcrumbs}</div>')
        if page_title:
            doc.append(f"<h1>{page_title}</h1>")
        doc.append(html_body)
        doc.append("</body></html>")

        ok, reason = render_html_to_pdf(
            "\n".join(doc), str(output_path), preference=self.engine_preference
        )
        if not ok:
            raise RuntimeError(f"PDF render failed: {reason}")


# ---------------------------------------------------------------------------
# Factory
# ---------------------------------------------------------------------------


FORMATS: dict[str, type[Formatter]] = {
    "html": HTMLFormatter,
    "md": MarkdownFormatter,
    "docx": DocxFormatter,
    "pdf": PDFFormatter,
}


def build_formatter(fmt: str, *, pdf_engine: str = "auto") -> Formatter:
    fmt = fmt.lower()
    cls = FORMATS.get(fmt)
    if cls is None:
        raise ValueError(
            f"Unknown format {fmt!r}. Choose from: {', '.join(FORMATS)}"
        )
    if cls is PDFFormatter:
        return PDFFormatter(engine=pdf_engine)
    return cls()
